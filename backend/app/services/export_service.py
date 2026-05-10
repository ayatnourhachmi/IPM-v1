"""Document export builders for delivery recommendations (PDF and DOCX)."""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

from app.core.recommendation_limits import (
    MAX_KPIS_RECOMMENDATIONS,
    MAX_ORGANIZATIONAL_RECOMMENDATIONS,
    MAX_TECHNICAL_RECOMMENDATIONS,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def _service_root() -> Path:
    return Path(__file__).resolve().parent.parent


def resolve_dxc_logo_path() -> Path | None:
    """Prefer bundled asset, then repo frontend landing asset (local dev)."""

    svc = Path(__file__).resolve()
    candidates = [
        _service_root() / "assets" / "dxc_logo.png",
        svc.parents[3] / "frontend" / "public" / "landing-files" / "DXC-Logo-2025.png",
    ]
    for p in candidates:
        if p.is_file():
            return p
    return None


def _fallback_logo_png_bytes() -> BytesIO:
    """Raster header when no official PNG is on disk."""
    from PIL import Image, ImageDraw, ImageFont

    w, h = 720, 112
    img = Image.new("RGB", (w, h), (8, 10, 14))
    draw = ImageDraw.Draw(img)
    for x in range(w):
        r = int(32 + min(140, x // 4))
        g = int(80 + min(120, x // 6))
        b = int(200 - min(140, x // 5))
        draw.line([(x, h - 10), (x, h - 4)], fill=(min(r, 255), min(g, 255), min(b, 255)))

    font_paths = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Bold.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    )
    sub_paths = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Regular.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    )
    font = None
    for fp in font_paths:
        try:
            font = ImageFont.truetype(fp, 34)
            break
        except OSError:
            continue
    if font is None:
        font = ImageFont.load_default()
    sub = None
    for fp in sub_paths:
        try:
            sub = ImageFont.truetype(fp, 15)
            break
        except OSError:
            continue
    if sub is None:
        sub = font

    draw.text((28, 20), "DXC Technology", fill=(250, 250, 252), font=font)
    draw.text((28, 64), "Innovation Progress Model · Delivery recommendations", fill=(180, 187, 196), font=sub)

    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _logo_flowable() -> RLImage:
    from PIL import Image as PILImage

    target_w = 2.65 * inch
    pth = resolve_dxc_logo_path()
    if pth is not None:
        im = PILImage.open(pth)
        tw, th = im.size
        h = target_w * th / tw
        return RLImage(str(pth), width=target_w, height=h)
    buf = _fallback_logo_png_bytes()
    im = PILImage.open(buf)
    buf.seek(0)
    tw, th = im.size
    h = target_w * th / tw
    return RLImage(ImageReader(buf), width=target_w, height=h)


def _organizational_line(item: object) -> str:
    if isinstance(item, dict):
        role = str(item.get("role", "") or "").strip()
        action = str(item.get("action", "") or "").strip()
        if role and action:
            return f"{role} — {action}"
        return action or role or "Not specified"
    return str(item)


def _plain_summary(pitch: str, n_solutions: int, n_recs: int) -> str:
    p = (pitch or "").strip()
    intro = (
        "This document summarises delivery recommendations from your innovation qualification process. "
        "It brings together technical themes, organisational alignment, and measurable outcomes for the "
        "selected solutions."
    )
    scope = (
        f"{n_recs} recommendation section(s) address {n_solutions} solution(s) included in this export. "
        "Governance, prioritisation, and commitment remain with your portfolio and executive teams."
    )
    if len(p) > 420:
        p = p[:417] + "…"
    return f"{intro} {scope} <b>Need snapshot:</b> {p or 'Not specified.'}"


_META_TABLE_CAPTION = (
    "<i>Left: field label. Right: value for this export (reference and date).</i>"
)

_DELIVERY_TABLE_CAPTION = (
    "<b>Reading this table:</b> <b>Solution</b> — name of the recommended offering. "
    "<b>Relevance</b> — estimated semantic fit between your stated need and the solution (0–100%). "
    "<b>Overall score</b> — composite evaluation score from the qualification step (scale used in the workshop, typically 1–5)."
)

_KPI_TABLE_CAPTION = (
    "<b>Reading this table:</b> <b>KPI</b> — outcome or indicator to track. "
    "<b>Target</b> — level or threshold to reach within the agreed horizon. "
    "<b>How we measure</b> — evidence, metric source, or governance review where success is verified."
)


def _word_count(blob: str) -> int:
    return len(re.findall(r"\w+", blob or "", flags=re.UNICODE))


def _pdf_table_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    """ReportLab Tables only wrap Flowables such as Paragraph; plain strings overlap."""
    s = str(text if text is not None else "").strip()
    return Paragraph(escape(s or "—"), style)


def _pdf_draw_footer(canvas, doc) -> None:  # noqa: ARG001
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    w, _ = A4
    canvas.drawCentredString(
        w / 2,
        34,
        "DXC Technology · Innovation Progress Model · For client / internal use under engagement terms",
    )
    n = canvas.getPageNumber()
    canvas.drawRightString(w - 42, 34, f"Page {n}")
    canvas.restoreState()


def build_pdf_report(
    need_id: str,
    pitch: str,
    recommendations: list[dict],
    delivery_solutions: list[dict],
) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=48,
        bottomMargin=56,
        title=f"IPM Recommendations - {need_id}",
        author="IPM / DXC",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#374151"),
    )
    subtle_style = ParagraphStyle(
        "Subtle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#6B7280"),
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#111827"),
        spaceBefore=14,
        spaceAfter=8,
    )
    subhead_style = ParagraphStyle(
        "Subhead",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1E3A5F"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14.5,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=4,
    )
    small_style = ParagraphStyle(
        "Small",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=8,
    )
    sol_tbl_hdr_para = ParagraphStyle(
        "SolTblHdrPara",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.white,
    )
    sol_tbl_cell_para = ParagraphStyle(
        "SolTblCellPara",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#1F2937"),
    )
    kpi_tbl_hdr_para = ParagraphStyle(
        "KpiTblHdrPara",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#1E1B4B"),
    )
    kpi_tbl_cell_para = ParagraphStyle(
        "KpiTblCellPara",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#1F2937"),
    )

    story: list = []
    story.append(_logo_flowable())
    story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            "Innovation Progress Model — <font color='#1E40AF'>Delivery Recommendations</font>",
            title_style,
        )
    )

    gen_ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    meta_rows = [
        ["Business need reference", need_id],
        ["Document generated", gen_ts],
        ["Recommendation blocks", str(len(recommendations))],
        ["Solutions in export", str(len(delivery_solutions))],
    ]
    meta_tbl = Table(meta_rows, colWidths=[150, 340])
    meta_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#334155")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E2E8F0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(meta_tbl)
    story.append(Paragraph(_META_TABLE_CAPTION, small_style))
    story.append(Spacer(1, 14))

    n_sol = len(delivery_solutions)
    summary = _plain_summary(pitch, n_sol, len(recommendations))
    story.append(Paragraph("<b>Executive summary</b>", subhead_style))
    story.append(Paragraph(summary, body_style))

    story.append(Paragraph("Business need (verbatim)", section_style))
    pitch_box = Table([[Paragraph(pitch or "Not specified", body_style)]], colWidths=[446])
    pitch_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story.append(pitch_box)
    story.append(
        Paragraph(
            f"<i>Approx. {_word_count(pitch)} words in the need statement above.</i>",
            small_style,
        )
    )
    story.append(Spacer(1, 10))

    if delivery_solutions:
        story.append(Paragraph("Selected delivery solutions", section_style))
        table_data = [
            [
                _pdf_table_paragraph("Solution", sol_tbl_hdr_para),
                _pdf_table_paragraph("Relevance", sol_tbl_hdr_para),
                _pdf_table_paragraph("Overall score", sol_tbl_hdr_para),
            ]
        ]
        for solution in delivery_solutions:
            table_data.append(
                [
                    _pdf_table_paragraph(solution.get("name", "Unknown"), sol_tbl_cell_para),
                    _pdf_table_paragraph(f"{solution.get('relevance', 0)}%", sol_tbl_cell_para),
                    _pdf_table_paragraph(f"{float(solution.get('overall', 0)):.2f}", sol_tbl_cell_para),
                ]
            )
        sol_tbl = Table(table_data, colWidths=[280, 90, 90])
        sol_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(sol_tbl)
        story.append(Paragraph(_DELIVERY_TABLE_CAPTION, small_style))
        story.append(Spacer(1, 12))

    for index, rec in enumerate(recommendations, start=1):
        if index > 1:
            story.append(PageBreak())

        title_line = rec.get("solution_name", "Solution")
        story.append(Paragraph(f"{index}. {title_line}", section_style))
        mode = str(rec.get("mode") or "STANDARD").upper()
        if mode == "PREREQUIS":
            story.append(
                Paragraph(
                    "<b>Recommendation mode: PREREQUIS</b> — solution fit is assessed as limited; "
                    "items below focus on readiness, prerequisites, and proofs—not a committed go-live plan.",
                    body_style,
                )
            )
            story.append(Spacer(1, 6))
        else:
            story.append(
                Paragraph(
                    "<b>Recommendation mode: STANDARD</b> — delivery-oriented guidance aligned with the "
                    "selected solution and your stated need.",
                    small_style,
                )
            )

        story.append(Paragraph("Technical recommendations", subhead_style))
        tech = rec.get("technical_recommendations", [])[:MAX_TECHNICAL_RECOMMENDATIONS]
        for i, item in enumerate(tech, start=1):
            story.append(Paragraph(f"{i}. {item}", body_style))
        if not tech:
            story.append(Paragraph("— None listed —", small_style))
        story.append(Spacer(1, 8))

        story.append(Paragraph("Organizational recommendations", subhead_style))
        org = rec.get("organizational_recommendations", [])[:MAX_ORGANIZATIONAL_RECOMMENDATIONS]
        for i, item in enumerate(org, start=1):
            story.append(Paragraph(f"{i}. {_organizational_line(item)}", body_style))
        if not org:
            story.append(Paragraph("— None listed —", small_style))
        story.append(Spacer(1, 8))

        story.append(Paragraph("Target KPIs and measurement criteria", subhead_style))
        kpis = rec.get("kpis", [])[:MAX_KPIS_RECOMMENDATIONS]
        if kpis:
            kpi_header = [
                [
                    _pdf_table_paragraph("KPI", kpi_tbl_hdr_para),
                    _pdf_table_paragraph("Target", kpi_tbl_hdr_para),
                    _pdf_table_paragraph("How we measure", kpi_tbl_hdr_para),
                ]
            ]
            kpi_rows = []
            for kpi in kpis:
                kpi_rows.append(
                    [
                        _pdf_table_paragraph(kpi.get("name", "KPI"), kpi_tbl_cell_para),
                        _pdf_table_paragraph(kpi.get("target", "—"), kpi_tbl_cell_para),
                        _pdf_table_paragraph(kpi.get("measurement_criteria", "—"), kpi_tbl_cell_para),
                    ]
                )
            # Extra width on Target avoids wrapped lines colliding with the next column (plain strings overflow).
            kt = Table(kpi_header + kpi_rows, colWidths=[100, 182, 164])
            kt.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2FF")),
                        ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#C7D2FE")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(kt)
            story.append(Paragraph(_KPI_TABLE_CAPTION, small_style))
        else:
            story.append(Paragraph("— None listed —", small_style))

        story.append(Spacer(1, 12))
        story.append(
            Paragraph(
                "Steer these recommendations through your portfolio / programme cadence; align owners, "
                "dates, and dependencies with enterprise architecture and sourcing before commitment.",
                small_style,
            )
        )

    doc.build(story, onFirstPage=_pdf_draw_footer, onLaterPages=_pdf_draw_footer)
    return buffer.getvalue()


def build_docx_report(
    need_id: str,
    pitch: str,
    recommendations: list[dict],
    delivery_solutions: list[dict],
) -> bytes:
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    logo_path = resolve_dxc_logo_path()
    if logo_path is not None:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Inches(2.6))
    else:
        buf = _fallback_logo_png_bytes()
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(buf, width=Inches(2.6))

    h1 = document.add_heading("Innovation Progress Model — Delivery recommendations", level=1)
    try:
        h1.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    except (IndexError, AttributeError):
        pass

    gen_ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    meta = document.add_paragraph()
    meta.add_run(f"Business need reference: {need_id}\n").bold = True
    meta.add_run(f"Generated: {gen_ts}\n")
    meta.add_run(f"Recommendation blocks: {len(recommendations)} | Solutions in export: {len(delivery_solutions)}")
    for r in meta.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    cap_meta = document.add_paragraph(
        "Left: field label. Right: value for this export (reference and date)."
    )
    cap_meta.runs[0].font.italic = True
    cap_meta.runs[0].font.size = Pt(9)
    cap_meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    document.add_heading("Executive summary", level=2)
    document.add_paragraph(
        _plain_summary(pitch, len(delivery_solutions), len(recommendations)).replace("<b>", "").replace("</b>", "")
    )

    document.add_heading("Business need (verbatim)", level=2)
    document.add_paragraph(pitch or "Not specified")
    wc_pitch = document.add_paragraph(
        f"Approx. {_word_count(pitch)} words in the need statement above."
    )
    wc_pitch.runs[0].font.italic = True

    if delivery_solutions:
        document.add_heading("Selected delivery solutions", level=2)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Solution"
        hdr[1].text = "Relevance"
        hdr[2].text = "Overall score"
        for solution in delivery_solutions:
            row = table.add_row().cells
            row[0].text = str(solution.get("name", "Unknown"))
            row[1].text = f"{solution.get('relevance', 0)}%"
            row[2].text = f"{float(solution.get('overall', 0)):.2f}"
        cap_del = document.add_paragraph()
        r0 = cap_del.add_run("Reading this table: ")
        r0.bold = True
        cap_del.add_run(
            "Solution — name of the recommended offering. "
            "Relevance — estimated semantic fit between your stated need and the solution (0–100%). "
            "Overall score — composite evaluation score from the qualification step "
            "(scale used in the workshop, typically 1–5)."
        )
        for r in cap_del.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        cap_del.runs[1].font.italic = True

    for index, rec in enumerate(recommendations, start=1):
        document.add_page_break()
        heading = f"{index}. {rec.get('solution_name', 'Solution')}"
        document.add_heading(heading, level=2)
        mode = str(rec.get("mode") or "STANDARD").upper()
        p_mode = document.add_paragraph()
        if mode == "PREREQUIS":
            r0 = p_mode.add_run("Recommendation mode: PREREQUIS — ")
            r0.bold = True
            p_mode.add_run(
                "solution fit is assessed as limited; items below focus on readiness, prerequisites, and "
                "proofs—not a committed go-live plan."
            )
        else:
            r0 = p_mode.add_run("Recommendation mode: STANDARD — ")
            r0.bold = True
            p_mode.add_run(
                "delivery-oriented guidance aligned with the selected solution and your stated need."
            )

        document.add_heading("Technical recommendations", level=3)
        tech = rec.get("technical_recommendations", [])[:MAX_TECHNICAL_RECOMMENDATIONS]
        for item in tech:
            document.add_paragraph(str(item), style="List Number")
        if not tech:
            document.add_paragraph("None listed.").runs[0].italic = True

        document.add_heading("Organizational recommendations", level=3)
        org = rec.get("organizational_recommendations", [])[:MAX_ORGANIZATIONAL_RECOMMENDATIONS]
        for item in org:
            document.add_paragraph(_organizational_line(item), style="List Number")
        if not org:
            document.add_paragraph("None listed.").runs[0].italic = True

        document.add_heading("Target KPIs and measurement criteria", level=3)
        kpis = rec.get("kpis", [])[:MAX_KPIS_RECOMMENDATIONS]
        if kpis:
            t = document.add_table(rows=1, cols=3)
            t.style = "Table Grid"
            h = t.rows[0].cells
            h[0].text = "KPI"
            h[1].text = "Target"
            h[2].text = "How we measure"
            for kpi in kpis:
                r = t.add_row().cells
                r[0].text = str(kpi.get("name", "KPI"))
                r[1].text = str(kpi.get("target", "—"))
                r[2].text = str(kpi.get("measurement_criteria", "—"))
            for row in t.rows:
                row.cells[0].width = Inches(1.15)
                row.cells[1].width = Inches(2.1)
                row.cells[2].width = Inches(2.2)
            cap_kpi = document.add_paragraph()
            k0 = cap_kpi.add_run("Reading this table: ")
            k0.bold = True
            cap_kpi.add_run(
                "KPI — outcome or indicator to track. "
                "Target — level or threshold to reach within the agreed horizon. "
                "How we measure — evidence, metric source, or governance review where success is verified."
            )
            for r in cap_kpi.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            cap_kpi.runs[1].font.italic = True
        else:
            document.add_paragraph("None listed.").runs[0].italic = True

        document.add_paragraph(
            "Steer these recommendations through your portfolio / programme cadence; align owners, "
            "dates, and dependencies with enterprise architecture and sourcing before commitment."
        )

    sec = document.sections[0]
    try:
        sec.footer.paragraphs[0].text = (
            "DXC Technology · Innovation Progress Model · For client / internal use under engagement terms"
        )
    except Exception:
        pass

    output = BytesIO()
    document.save(output)
    return output.getvalue()
